from docx import Document
from datetime import datetime
import os
import sqlite3

LOG_DB_PATH = "usb_logs.db"
REPORT_DIR = "reports"

def generate_word_report(date: datetime.date):
    date_str = date.strftime("%Y-%m-%d")

    os.makedirs(REPORT_DIR, exist_ok=True)
    report_path = os.path.join(REPORT_DIR, f"Отчет_{date.strftime('%d-%m-%Y')}.docx")


    doc = Document()
    doc.add_heading(f"Отчет за {date.strftime('%d.%m.%Y')}", level=1)

    conn = sqlite3.connect(LOG_DB_PATH)
    cursor = conn.cursor()
    cursor.execute("""
        SELECT timestamp, file_name, src_path, dst_path, action, status 
        FROM file_logs 
        WHERE DATE(timestamp) = ?
        ORDER BY timestamp ASC
    """, (date_str,))
    rows = cursor.fetchall()
    conn.close()

    if not rows:
        doc.add_paragraph("Нет записей за выбранную дату.")
    else:
        for row in rows:
            timestamp, file_name, src_path, dst_path, action, status = row
            doc.add_paragraph(f"[{timestamp}] {action} | {file_name}\nИз: {src_path}\nВ:  {dst_path}\nСтатус: {status}")

    doc.save(report_path)
    return report_path
